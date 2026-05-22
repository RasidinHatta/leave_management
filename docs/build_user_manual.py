from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "HR_Leave_Management_User_Manual.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(31, 41, 55)
MUTED = RGBColor(91, 103, 115)
FILL = "E8EEF5"
LIGHT_FILL = "F4F6F9"
BORDER = "B8C2CC"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_grid = table._tbl.tblGrid
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_borders(table, color=BORDER):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def style_run(run, bold=False, italic=False, color=INK, size=11):
    run.bold = bold
    run.italic = italic
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.color.rgb = color


def add_title(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("HR Leave Management")
    style_run(r, bold=True, color=RGBColor(11, 37, 69), size=22)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run("User Manual and Operating Guide")
    style_run(r, color=MUTED, size=12)

    table = doc.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [2200, 7160])
    set_borders(table)
    rows = [
        ("Application", "Windows desktop leave administration tool"),
        ("Primary users", "ADMIN, USER, and REPORT roles"),
        ("Configuration", "SQL Server connections are read from config.ini"),
    ]
    for row, (label, detail) in zip(table.rows, rows):
        set_cell_shading(row.cells[0], FILL)
        row.cells[0].paragraphs[0].add_run(label).bold = True
        row.cells[1].paragraphs[0].add_run(detail)
    doc.add_paragraph()


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    style_run(run)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        style_run(run)


def add_steps(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        style_run(run)


def add_note(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [9360])
    set_borders(table, "D8DEE8")
    set_cell_shading(table.cell(0, 0), LIGHT_FILL)
    p = table.cell(0, 0).paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    style_run(r, color=DARK_BLUE)
    doc.add_paragraph()


def add_role_table(doc):
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    set_table_geometry(table, [1800, 3480, 4080])
    set_borders(table)
    headers = ["Role", "Visible menus", "Purpose"]
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, FILL)
        p = cell.paragraphs[0]
        r = p.add_run(header)
        style_run(r, bold=True, color=DARK_BLUE)
    rows = [
        ("ADMIN", "Bring Forward, Leave Taken, DB Targets, Leave Report Config, Manage Users", "Full application administration."),
        ("USER", "Bring Forward, Leave Taken", "Daily leave data entry and upload work."),
        ("REPORT", "Leave Report Config", "Maintain daily report targets only."),
    ]
    for role, menus, purpose in rows:
        row = table.add_row().cells
        for idx, value in enumerate((role, menus, purpose)):
            p = row[idx].paragraphs[0]
            r = p.add_run(value)
            style_run(r, bold=(idx == 0))
    doc.add_paragraph()


def add_config_table(doc):
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    set_table_geometry(table, [2100, 2500, 4760])
    set_borders(table)
    headers = ["Section", "Key", "Description"]
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, FILL)
        r = cell.paragraphs[0].add_run(header)
        style_run(r, bold=True, color=DARK_BLUE)
    rows = [
        ("DatabaseConfig", "Server", "Main SQL Server instance for leave operations."),
        ("DatabaseConfig", "Database", "Main application database, for example MYPAY_LCO."),
        ("DatabaseConfig", "Driver", "ODBC driver name, usually ODBC Driver 17 for SQL Server."),
        ("ReportConfig", "Server", "SQL Server for HR_REPORT_CONFIG."),
        ("ReportConfig", "Driver", "ODBC driver for HR_REPORT_CONFIG."),
    ]
    for values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(values):
            r = cells[idx].paragraphs[0].add_run(value)
            style_run(r)
    doc.add_paragraph()


def setup_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def build():
    doc = Document()
    setup_styles(doc)
    add_title(doc)

    add_heading(doc, "1. Overview", 1)
    add_body(
        doc,
        "HR Leave Management is a Windows desktop application for managing carry-forward leave, leave taken entries, report target configuration, and user access. It connects directly to SQL Server using the settings in config.ini.",
    )
    add_note(
        doc,
        "Important: Keep config.ini beside leave_management.exe in the Release folder. The application reads that file first so production database settings can be changed without rebuilding the app.",
    )

    add_heading(doc, "2. User Roles and Menu Access", 1)
    add_role_table(doc)

    add_heading(doc, "3. Sign In and General Navigation", 1)
    add_steps(
        doc,
        [
            "Open leave_management.exe.",
            "Enter the username and password created in the application database.",
            "Use the left sidebar to open the menus available to your role.",
            "Use the settings icon in the sidebar footer to change font size, display mode, or color theme.",
            "Use the logout icon in the sidebar footer when work is complete.",
        ],
    )

    add_heading(doc, "4. Bring Forward Leave", 1)
    add_body(doc, "Use Bring Forward to add carry-forward annual leave records.")
    add_steps(
        doc,
        [
            "Confirm the displayed database is correct.",
            "Select Target Year and Target Month.",
            "Enter Employee Code and Days Leave manually, or use Import Excel.",
            "Check the valid row count before submitting.",
            "Click Run Bring Forward to execute the database operation.",
        ],
    )
    add_bullets(
        doc,
        [
            "Export Excel creates a template for bulk entry.",
            "Imported rows are validated before submission.",
            "Invalid rows are highlighted and must be corrected before running.",
        ],
    )

    add_heading(doc, "5. Leave Taken", 1)
    add_body(doc, "Use Leave Taken to add employee leave records for specific dates and leave types.")
    add_steps(
        doc,
        [
            "Confirm the displayed database is correct.",
            "Enter Employee Code and Leave Date in YYYY-MM-DD format.",
            "Choose the Leave Type loaded from dbo.LV_TYPE.",
            "Enter an optional remark.",
            "Click Run Leave Taken to write valid rows to SQL Server.",
        ],
    )

    add_heading(doc, "6. DB Targets", 1)
    add_body(
        doc,
        "DB Targets shows the main database configuration loaded from config.ini. Use Test Connection to verify that the configured SQL Server, database, and ODBC driver are reachable.",
    )

    add_heading(doc, "7. Leave Report Config", 1)
    add_body(
        doc,
        "Leave Report Config is available to ADMIN and REPORT users. It connects to the fixed database HR_REPORT_CONFIG using the ReportConfig section in config.ini.",
    )
    add_steps(
        doc,
        [
            "Click Test Connection to confirm the report configuration server is reachable.",
            "Click Setup DB if HR_REPORT_CONFIG or dbo.report_targets must be created or repaired.",
            "Click Add Target to create a report target.",
            "Use Edit to update SMTP settings, recipient lists, TLS, and active status.",
            "Use Delete only when the report target is no longer required.",
        ],
    )
    add_note(
        doc,
        "Email passwords are stored in the email_password column and encrypted in SQL Server. When editing a target, leave the password field blank to keep the current password.",
    )

    add_heading(doc, "8. Manage Users", 1)
    add_body(doc, "Manage Users is available to ADMIN users only.")
    add_bullets(
        doc,
        [
            "ADMIN accounts are hidden in the list and cannot be created or edited from this screen.",
            "New users can be created only as USER or REPORT.",
            "USER accounts can access the main leave menus.",
            "REPORT accounts can access only Leave Report Config.",
        ],
    )

    add_heading(doc, "9. Configuration Reference", 1)
    add_config_table(doc)

    add_heading(doc, "10. Troubleshooting", 1)
    add_bullets(
        doc,
        [
            "Connection failure: confirm Server, Database, and Driver in config.ini and verify the ODBC driver is installed.",
            "Memory allocation failure while reading data: retry after updating to the latest release build; the app uses bounded SQL casts for known ODBC buffer issues.",
            "Leave types do not load: verify dbo.LV_TYPE exists and contains rows where LV_EVENT_CODE equals LEAVE.",
            "Report targets table missing: open Leave Report Config and click Setup DB.",
            "Email sending fails: verify SMTP server, port, user, password, TLS setting, and provider app-password requirements.",
        ],
    )

    add_heading(doc, "11. Safe Operating Checklist", 1)
    add_bullets(
        doc,
        [
            "Back up SQL databases before first production use.",
            "Keep config.ini with the release executable.",
            "Restrict ADMIN credentials to trusted operators.",
            "Test SQL and report configuration connections after changing config.ini.",
            "Review imported Excel data before running bulk operations.",
        ],
    )

    doc.save(OUT)


if __name__ == "__main__":
    build()
